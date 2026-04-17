"""SEO Audit Team router — all routes live under /api/seo-audit/"""

import asyncio
import io
import os
import uuid

from docx import Document
from fastapi import APIRouter, Cookie, Depends
from fastapi.responses import JSONResponse, Response, StreamingResponse
from pydantic import BaseModel

import sys

from state import get_session, save_session, get_user, get_user_by_email, log_activity, get_token_email, log_history_item, get_rollback_stage, redis_client
from utils.usage import ToolAccess, increment_usage
from utils.sanitise import sanitise_user_input, validate_url
from utils.sse import SSE_HEADERS, sse_chunk, sse_done, sse_event, friendly_error
from agents.seo_audit import auditor, analyser, recommender, implementer
from services.agency_log import log_task
from services.notion_seo_audit import save_audit_report

router = APIRouter()

_SESSION_DEFAULTS = {
    "stage": "idle",
    "url": "",
    "audit_context": "",
    "competitor_urls": [],
    "audit_data": {},
    "audit_text": "",
    "analysis": "",
    "recommendations": "",
    "implementation": "",
    "notion_url": None,
    "user_id": "",
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

async def _get_api_key(session: dict, agency_token: str | None = None) -> str:
    user = await get_user(session.get("user_id", ""))
    key = (user or {}).get("gemini_api_key") or os.environ.get("GEMINI_API_KEY", "")
    if not key and agency_token:
        email = await get_token_email(agency_token)
        if email:
            user = await get_user_by_email(email)
            key = (user or {}).get("gemini_api_key") or ""
    return key


async def _get_notion_creds(session: dict) -> tuple[str, str]:
    user = await get_user(session.get("user_id", ""))
    u = user or {}
    return u.get("notion_token", ""), u.get("notion_on_page_db_id", "") or u.get("notion_seo_audit_db_id", "")


# ---------------------------------------------------------------------------
# Pydantic models
# ---------------------------------------------------------------------------

class CreateSessionPayload(BaseModel):
    user_id: str = ""


class StartAuditRequest(BaseModel):
    session_id: str
    url: str
    context: str
    competitor_urls: list[str] = []


class SessionRequest(BaseModel):
    session_id: str


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@router.post("/session")
async def create_session(
    payload: CreateSessionPayload,
    _: None = Depends(ToolAccess("seo_audit")),
):
    sid = str(uuid.uuid4())
    defaults = {**_SESSION_DEFAULTS, "user_id": payload.user_id}
    await get_session(sid, "seo_audit", defaults)
    return {"session_id": sid}


@router.get("/state")
async def get_state(session_id: str):
    sess = await get_session(session_id, "seo_audit", _SESSION_DEFAULTS)
    return JSONResponse(sess)


_STUCK_STAGES = {"auditing", "analysing", "recommending", "implementing"}

@router.post("/start")
async def start_audit(req: StartAuditRequest, agency_token: str | None = Cookie(default=None)):
    sess = await get_session(req.session_id, "seo_audit", _SESSION_DEFAULTS)
    if sess["stage"] in _STUCK_STAGES:
        # SSE connection likely dropped (common on iPad/Safari) — auto-reset and continue
        sess["stage"] = "idle"
    if sess["stage"] not in ("idle", "done"):
        return JSONResponse({"error": "Audit already in progress"}, status_code=400)
    sess["url"] = validate_url(req.url)
    sess["audit_context"] = sanitise_user_input(req.context, user_id=sess.get("user_id"))
    sess["competitor_urls"] = [validate_url(u.strip()) for u in req.competitor_urls if u.strip()]
    sess["stage"] = "auditing"
    sess["audit_data"] = {}
    sess["audit_text"] = ""
    sess["analysis"] = ""
    sess["recommendations"] = ""
    sess["implementation"] = ""
    sess["notion_url"] = None
    await save_session(req.session_id, sess)
    email = await get_token_email(agency_token) if agency_token else None
    await log_activity("seo_audit", f"Started audit: {req.url}", email=email)
    return {"ok": True}


@router.get("/stream/audit")
async def stream_audit(session_id: str, agency_token: str | None = Cookie(default=None)):
    sess = await get_session(session_id, "seo_audit", _SESSION_DEFAULTS)
    if sess["stage"] != "auditing":
        return JSONResponse({"error": "Not in auditing stage"}, status_code=400)

    api_key = await _get_api_key(sess, agency_token)
    if not api_key:
        async def no_key(): yield sse_event({"type": "error", "code": "gemini_not_configured", "message": "Gemini API key not set. Open Settings to configure it."})
        return StreamingResponse(no_key(), media_type="text/event-stream", headers=SSE_HEADERS)

    email = await get_token_email(agency_token) if agency_token else None

    async def generate():
        full_text = ""
        try:
            async for kind, value in auditor.run(sess["url"], sess["audit_context"], api_key=api_key):
                if kind == "technical_signals":
                    yield sse_event({"type": "technical_signals", "data": value})
                elif kind == "chunk":
                    full_text += value
                    yield sse_chunk(value)
                elif kind == "audit_data":
                    sess["audit_data"] = value
                    sess["audit_text"] = full_text
                    sess["stage"] = "awaiting_analyse"
                    await save_session(session_id, sess)
                    yield sse_event({"type": "audit_data", "data": value})
                    if email:
                        await increment_usage(redis_client, email, "seo_audit")
                    yield sse_done()
                    return
                elif kind == "error":
                    sess["stage"] = get_rollback_stage("seo_audit", "auditing")
                    await save_session(session_id, sess)
                    yield sse_event({"type": "error", "message": friendly_error(value)})
                    yield sse_done()
                    return
        except Exception as exc:
            print(f"[ERROR] seo_audit/audit session={session_id}: {exc}", file=sys.stderr)
            sess["stage"] = get_rollback_stage("seo_audit", "auditing")
            await save_session(session_id, sess)
            yield sse_event({"type": "error", "message": "Agent failed. Please try again."})
            yield sse_done()

    return StreamingResponse(generate(), media_type="text/event-stream", headers=SSE_HEADERS)


@router.post("/analyse")
async def start_analyse(req: SessionRequest):
    sess = await get_session(req.session_id, "seo_audit", _SESSION_DEFAULTS)
    if sess["stage"] != "awaiting_analyse":
        return JSONResponse({"error": "Not ready to analyse"}, status_code=400)
    sess["stage"] = "analysing"
    await save_session(req.session_id, sess)
    return {"ok": True}


@router.get("/stream/analysis")
async def stream_analysis(session_id: str, agency_token: str | None = Cookie(default=None)):
    sess = await get_session(session_id, "seo_audit", _SESSION_DEFAULTS)
    if sess["stage"] != "analysing":
        return JSONResponse({"error": "Not in analysing stage"}, status_code=400)

    api_key = await _get_api_key(sess, agency_token)
    if not api_key:
        async def no_key(): yield sse_event({"type": "error", "code": "gemini_not_configured", "message": "Gemini API key not set. Open Settings to configure it."})
        return StreamingResponse(no_key(), media_type="text/event-stream", headers=SSE_HEADERS)

    async def generate():
        full_text = ""
        try:
            async for kind, value in analyser.run(
                sess["url"], sess["audit_context"], sess["audit_data"],
                api_key=api_key, competitor_urls=sess.get("competitor_urls", [])
            ):
                if kind == "chunk":
                    full_text += value
                    yield sse_chunk(value)
                elif kind == "done":
                    sess["analysis"] = full_text
                    sess["stage"] = "awaiting_recommend"
                    await save_session(session_id, sess)
                    yield sse_done()
                    return
                elif kind == "error":
                    sess["stage"] = get_rollback_stage("seo_audit", "analysing")
                    await save_session(session_id, sess)
                    yield sse_event({"type": "error", "message": friendly_error(value)})
                    yield sse_done()
                    return
        except Exception as exc:
            print(f"[ERROR] seo_audit/analysis session={session_id}: {exc}", file=sys.stderr)
            sess["stage"] = get_rollback_stage("seo_audit", "analysing")
            await save_session(session_id, sess)
            yield sse_event({"type": "error", "message": "Agent failed. Please try again."})
            yield sse_done()

    return StreamingResponse(generate(), media_type="text/event-stream", headers=SSE_HEADERS)


@router.post("/recommend")
async def start_recommend(req: SessionRequest):
    sess = await get_session(req.session_id, "seo_audit", _SESSION_DEFAULTS)
    if sess["stage"] != "awaiting_recommend":
        return JSONResponse({"error": "Not ready to recommend"}, status_code=400)
    sess["stage"] = "recommending"
    await save_session(req.session_id, sess)
    return {"ok": True}


@router.get("/stream/recommendations")
async def stream_recommendations(session_id: str, agency_token: str | None = Cookie(default=None)):
    sess = await get_session(session_id, "seo_audit", _SESSION_DEFAULTS)
    if sess["stage"] != "recommending":
        return JSONResponse({"error": "Not in recommending stage"}, status_code=400)

    api_key = await _get_api_key(sess, agency_token)
    if not api_key:
        async def no_key(): yield sse_event({"type": "error", "code": "gemini_not_configured", "message": "Gemini API key not set. Open Settings to configure it."})
        return StreamingResponse(no_key(), media_type="text/event-stream", headers=SSE_HEADERS)

    async def generate():
        full_text = ""
        try:
            async for kind, value in recommender.run(
                sess["url"], sess["audit_context"], sess["audit_data"], sess["analysis"],
                api_key=api_key, competitor_urls=sess.get("competitor_urls", [])
            ):
                if kind == "chunk":
                    full_text += value
                    yield sse_chunk(value)
                elif kind == "done":
                    sess["recommendations"] = full_text
                    sess["stage"] = "awaiting_implement"
                    await save_session(session_id, sess)
                    yield sse_done()
                    return
                elif kind == "error":
                    sess["stage"] = get_rollback_stage("seo_audit", "recommending")
                    await save_session(session_id, sess)
                    yield sse_event({"type": "error", "message": friendly_error(value)})
                    yield sse_done()
                    return
        except Exception as exc:
            print(f"[ERROR] seo_audit/recommendations session={session_id}: {exc}", file=sys.stderr)
            sess["stage"] = get_rollback_stage("seo_audit", "recommending")
            await save_session(session_id, sess)
            yield sse_event({"type": "error", "message": "Agent failed. Please try again."})
            yield sse_done()

    return StreamingResponse(generate(), media_type="text/event-stream", headers=SSE_HEADERS)


@router.post("/implement")
async def start_implement(req: SessionRequest):
    sess = await get_session(req.session_id, "seo_audit", _SESSION_DEFAULTS)
    if sess["stage"] != "awaiting_implement":
        return JSONResponse({"error": "Not ready to implement"}, status_code=400)
    sess["stage"] = "implementing"
    await save_session(req.session_id, sess)
    return {"ok": True}


@router.get("/stream/implementation")
async def stream_implementation(session_id: str, agency_token: str | None = Cookie(default=None)):
    sess = await get_session(session_id, "seo_audit", _SESSION_DEFAULTS)
    if sess["stage"] != "implementing":
        return JSONResponse({"error": "Not in implementing stage"}, status_code=400)

    api_key = await _get_api_key(sess, agency_token)
    if not api_key:
        async def no_key(): yield sse_event({"type": "error", "code": "gemini_not_configured", "message": "Gemini API key not set. Open Settings to configure it."})
        return StreamingResponse(no_key(), media_type="text/event-stream", headers=SSE_HEADERS)
    email = await get_token_email(agency_token) if agency_token else None

    async def generate():
        full_text = ""
        cms = sess["audit_data"].get("cms", "WordPress")
        try:
            async for kind, value in implementer.run(
                sess["url"],
                sess["audit_context"],
                cms,
                sess["audit_data"],
                sess["analysis"],
                sess["recommendations"],
                api_key=api_key,
            ):
                if kind == "chunk":
                    full_text += value
                    yield sse_chunk(value)
                elif kind == "done":
                    sess["implementation"] = full_text
                    sess["stage"] = "done"
                    await save_session(session_id, sess)
                    await log_activity("seo_audit", f"Completed audit: {sess['url']}", email=email)
                    user = await get_user(sess.get("user_id", ""))
                    u = user or {}
                    asyncio.ensure_future(log_task(
                        "SEO Audit", "SEO Audit", f"Audit: {sess['url']}",
                        notion_token=u.get("notion_token", ""),
                        db_id=u.get("notion_agency_log_db_id", ""),
                    ))
                    if email and full_text:
                        await log_history_item(email, "SEO Audit", f"Audit: {sess['url']}", full_text)
                    yield sse_done()
                    return
                elif kind == "error":
                    sess["stage"] = get_rollback_stage("seo_audit", "implementing")
                    await save_session(session_id, sess)
                    yield sse_event({"type": "error", "message": friendly_error(value)})
                    yield sse_done()
                    return
        except Exception as exc:
            print(f"[ERROR] seo_audit/implementation session={session_id}: {exc}", file=sys.stderr)
            sess["stage"] = get_rollback_stage("seo_audit", "implementing")
            await save_session(session_id, sess)
            yield sse_event({"type": "error", "message": "Agent failed. Please try again."})
            yield sse_done()

    return StreamingResponse(generate(), media_type="text/event-stream", headers=SSE_HEADERS)


@router.post("/save-to-notion")
async def save_to_notion(req: SessionRequest):
    sess = await get_session(req.session_id, "seo_audit", _SESSION_DEFAULTS)
    if sess["stage"] != "done":
        return JSONResponse({"error": "Audit not complete"}, status_code=400)
    if sess.get("notion_url"):
        return {"notion_url": sess["notion_url"]}

    user = await get_user(sess.get("user_id", ""))
    u = user or {}
    notion_token = u.get("notion_token", "")
    db_id = u.get("notion_seo_audit_db_id", "")

    if not notion_token or not db_id:
        return JSONResponse(
            {"error": "Notion saving isn't configured for your account yet.", "code": "notion_not_configured"},
            status_code=400,
        )

    try:
        notion_url = await save_audit_report(
            url=sess["url"],
            audit_data=sess["audit_data"],
            audit_text=sess["audit_text"],
            analysis=sess["analysis"],
            recommendations=sess["recommendations"],
            implementation=sess["implementation"],
            notion_token=notion_token,
            db_id=db_id,
        )
        sess["notion_url"] = notion_url
        await save_session(req.session_id, sess)
        if notion_url:
            asyncio.ensure_future(log_task(
                "SEO Audit", "SEO Audit", f"Saved to Notion: {sess['url']}", link=notion_url,
                notion_token=u.get("notion_token", ""),
                db_id=u.get("notion_agency_log_db_id", ""),
            ))
        return {"notion_url": notion_url}
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)


@router.get("/download")
async def download_report(session_id: str):
    sess = await get_session(session_id, "seo_audit", _SESSION_DEFAULTS)
    if sess["stage"] != "done":
        return JSONResponse({"error": "Audit not complete"}, status_code=400)

    url = sess.get("url", "Unknown URL")
    doc = Document()
    doc.add_heading(f"SEO Audit Report — {url}", level=1)

    sections = [
        ("Audit Findings", sess.get("audit_text", "")),
        ("Analysis", sess.get("analysis", "")),
        ("Recommendations", sess.get("recommendations", "")),
        ("Implementation Guide", sess.get("implementation", "")),
    ]
    for heading, text in sections:
        if text:
            doc.add_heading(heading, level=2)
            for line in text.splitlines():
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_url = "".join(c if c.isalnum() or c in "-_." else "_" for c in url.replace("https://", "").replace("http://", ""))[:50]
    filename = f"SEO_Audit_{safe_url}.docx"

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/reset")
async def reset_audit(req: SessionRequest):
    sess = await get_session(req.session_id, "seo_audit", _SESSION_DEFAULTS)
    sess.update({
        "url": "",
        "audit_context": "",
        "competitor_urls": [],
        "stage": "idle",
        "audit_data": {},
        "audit_text": "",
        "analysis": "",
        "recommendations": "",
        "implementation": "",
        "notion_url": None,
    })
    await save_session(req.session_id, sess)
    return {"ok": True}
